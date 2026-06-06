"""Word (.docx) parser using python-docx heading styles."""

import io
import logging

from src.parsers.base import DocumentSection

logger = logging.getLogger(__name__)

# Mapping from python-docx style names to heading levels
HEADING_STYLE_MAP = {
    "Heading 1": 1,
    "Heading 2": 2,
    "Heading 3": 3,
    "Heading 4": 3,
    "Title": 1,
}


class DocxParser:
    """Extract structured sections from Word documents."""

    def parse(self, content: bytes) -> list[DocumentSection]:
        from docx import Document

        doc = Document(io.BytesIO(content))
        sections: list[DocumentSection] = []
        current_title = ""
        current_level = 1
        current_lines: list[str] = []
        chapter_title = ""

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue

            style_name = para.style.name if para.style else ""
            level = HEADING_STYLE_MAP.get(style_name)

            if level is not None:
                # Flush previous section
                if current_lines:
                    sections.append(
                        DocumentSection(
                            title=current_title,
                            level=current_level,
                            content="\n".join(current_lines),
                            parent_title=chapter_title if current_level > 1 else "",
                        )
                    )
                    current_lines = []

                current_title = text
                current_level = level
                if level == 1:
                    chapter_title = text
            else:
                current_lines.append(text)

        # Flush last section
        if current_lines:
            sections.append(
                DocumentSection(
                    title=current_title,
                    level=current_level,
                    content="\n".join(current_lines),
                    parent_title=chapter_title if current_level > 1 else "",
                )
            )

        # Fallback: if no headings detected, return entire doc as one section
        if not sections:
            all_text = "\n".join(
                p.text.strip() for p in doc.paragraphs if p.text.strip()
            )
            if all_text:
                sections.append(
                    DocumentSection(title="", level=1, content=all_text)
                )

        logger.info("[DocxParser] extracted %d sections", len(sections))
        return sections
